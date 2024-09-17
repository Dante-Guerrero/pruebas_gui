import tkinter as tk
from tkinter import messagebox, ttk
from tkcalendar import DateEntry
from fpdf import FPDF
from datetime import datetime
from docx import Document
import pandas as pd

def calculate_age(birthdate):
    today = datetime.today()
    age = today.year - birthdate.year - ((today.month, today.day) < (birthdate.month, birthdate.day))
    return age

def create_pdf():
    # Obtener los datos de las entradas
    nombre = entry_nombre.get()
    apellido = entry_apellido.get()
    birthdate = date_entry.get_date()
    sexo = var_sexo.get()
    profesion = combo_profesion.get()
    talla = spin_talla.get()
    habilidades = [habilidad for habilidad, var in habilidades_vars.items() if var.get()]
    
    if not nombre or not apellido or not birthdate or not sexo or not profesion or not talla:
        messagebox.showwarning("Advertencia", "Todos los campos son obligatorios.")
        return
    
    edad = calculate_age(birthdate)
    
    # Crear un objeto PDF
    pdf = FPDF()
    pdf.add_page()
    
    # Establecer la fuente
    pdf.set_font("Arial", size=12)
    
    # Añadir texto al PDF
    pdf.cell(200, 10, txt=f"Nombre: {nombre}", ln=True)
    pdf.cell(200, 10, txt=f"Apellido: {apellido}", ln=True)
    pdf.cell(200, 10, txt=f"Fecha de Nacimiento: {birthdate.strftime('%d/%m/%Y')}", ln=True)
    pdf.cell(200, 10, txt=f"Edad: {edad}", ln=True)
    pdf.cell(200, 10, txt=f"Sexo: {sexo}", ln=True)
    pdf.cell(200, 10, txt=f"Profesión: {profesion}", ln=True)
    pdf.cell(200, 10, txt=f"Talla: {talla}", ln=True)
    pdf.cell(200, 10, txt="Habilidades:", ln=True)
    
    for habilidad in habilidades:
        pdf.cell(200, 10, txt=f"- {habilidad}", ln=True)
    
    # Guardar el PDF
    pdf.output("datos_registrados.pdf")
    
    messagebox.showinfo("Éxito", "El PDF ha sido creado exitosamente.")

def create_word():
    # Obtener los datos de las entradas
    nombre = entry_nombre.get()
    apellido = entry_apellido.get()
    birthdate = date_entry.get_date()
    sexo = var_sexo.get()
    profesion = combo_profesion.get()
    talla = spin_talla.get()
    habilidades = [habilidad for habilidad, var in habilidades_vars.items() if var.get()]
    
    if not nombre or not apellido or not birthdate or not sexo or not profesion or not talla:
        messagebox.showwarning("Advertencia", "Todos los campos son obligatorios.")
        return
    
    edad = calculate_age(birthdate)
    
    # Crear un documento Word
    doc = Document()
    
    # Añadir texto al documento
    doc.add_heading('Datos Registrados', level=1)
    doc.add_paragraph(f"Nombre: {nombre}")
    doc.add_paragraph(f"Apellido: {apellido}")
    doc.add_paragraph(f"Fecha de Nacimiento: {birthdate.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Edad: {edad}")
    doc.add_paragraph(f"Sexo: {sexo}")
    doc.add_paragraph(f"Profesión: {profesion}")
    doc.add_paragraph(f"Talla: {talla}")
    
    doc.add_paragraph("Habilidades:")
    
    for habilidad in habilidades:
        doc.add_paragraph(f"- {habilidad}")
    
    # Guardar el documento
    doc.save('datos_registrados.docx')
    
    messagebox.showinfo("Éxito", "El documento Word ha sido creado exitosamente.")

def create_excel():
    # Obtener los datos de las entradas
    nombre = entry_nombre.get()
    apellido = entry_apellido.get()
    birthdate = date_entry.get_date()
    sexo = var_sexo.get()
    profesion = combo_profesion.get()
    talla = spin_talla.get()
    habilidades = [habilidad for habilidad, var in habilidades_vars.items() if var.get()]
    
    if not nombre or not apellido or not birthdate or not sexo or not profesion or not talla:
        messagebox.showwarning("Advertencia", "Todos los campos son obligatorios.")
        return
    
    edad = calculate_age(birthdate)
    
    # Crear un DataFrame de pandas
    data = {
        'Nombre': [nombre],
        'Apellido': [apellido],
        'Fecha de Nacimiento': [birthdate.strftime('%d/%m/%Y')],
        'Edad': [edad],
        'Sexo': [sexo],
        'Profesión': [profesion],
        'Talla': [talla],
        'Habilidades': [', '.join(habilidades)]
    }
    
    df = pd.DataFrame(data)
    
    # Guardar el DataFrame en un archivo Excel
    df.to_excel('datos_registrados.xlsx', index=False)
    
    messagebox.showinfo("Éxito", "El archivo Excel ha sido creado exitosamente.")

def update_age(event):
    birthdate = date_entry.get_date()
    edad = calculate_age(birthdate)
    label_edad.config(text=f"Edad: {edad}")

# Crear la ventana principal
root = tk.Tk()
root.title("Formulario de Prueba")

# Establecer el tamaño de la ventana
root.geometry("500x800")

# Establecer el color de fondo
root.configure(bg="#f0f0f0")

# Crear un título para el formulario
title = tk.Label(root, text="FORMULARIO DE PRUEBA", font=("Helvetica", 24, "bold"), bg="#f0f0f0", fg="#333333")
title.pack(pady=20)

# Crear un marco para el formulario
frame = tk.Frame(root, bg="#ffffff", padx=20, pady=20, relief=tk.RAISED, bd=2)
frame.pack(pady=20)

# Crear y colocar las etiquetas y entradas para nombre y apellido
tk.Label(frame, text="Nombre", bg="#ffffff", font=("Helvetica", 12)).grid(row=0, column=0, pady=5, sticky="e")
entry_nombre = tk.Entry(frame, font=("Helvetica", 12))
entry_nombre.grid(row=0, column=1, pady=5)

tk.Label(frame, text="Apellido", bg="#ffffff", font=("Helvetica", 12)).grid(row=1, column=0, pady=5, sticky="e")
entry_apellido = tk.Entry(frame, font=("Helvetica", 12))
entry_apellido.grid(row=1, column=1, pady=5)

# Crear y colocar los radiobuttons para seleccionar el sexo
tk.Label(frame, text="Sexo", bg="#ffffff", font=("Helvetica", 12)).grid(row=2, column=0, pady=5, sticky="e")
var_sexo = tk.StringVar(value="Masculino")
tk.Radiobutton(frame, text="Masculino", variable=var_sexo, value="Masculino", bg="#ffffff", font=("Helvetica", 12)).grid(row=2, column=1, sticky="w")
tk.Radiobutton(frame, text="Femenino", variable=var_sexo, value="Femenino", bg="#ffffff", font=("Helvetica", 12)).grid(row=2, column=2, sticky="w")

# Crear y colocar el DateEntry para la fecha de nacimiento
tk.Label(frame, text="Fecha de Nacimiento", bg="#ffffff", font=("Helvetica", 12)).grid(row=3, column=0, pady=5, sticky="e")
date_entry = DateEntry(frame, selectmode='day', year=2000, month=1, day=1, font=("Helvetica", 12))
date_entry.grid(row=3, column=1, pady=5)
date_entry.bind("<<DateEntrySelected>>", update_age)

# Crear y colocar la etiqueta para mostrar la edad calculada
label_edad = tk.Label(frame, text="Edad: ", bg="#ffffff", font=("Helvetica", 12))
label_edad.grid(row=4, columnspan=2)

# Crear y colocar el combobox para seleccionar la profesión
tk.Label(frame, text="Profesión", bg="#ffffff", font=("Helvetica", 12)).grid(row=5, column=0, pady=5, sticky="e")
profesiones = ["Abogado", "Economista", "Ingeniero", "Médico", "Profesor", "Arquitecto", "Diseñador", "Científico", "Artista", "Escritor"]
combo_profesion = ttk.Combobox(frame, values=profesiones, font=("Helvetica", 12))
combo_profesion.grid(row=5, column=1, pady=5)

# Crear y colocar el spinbox para ingresar la talla
tk.Label(frame, text="Talla", bg="#ffffff", font=("Helvetica", 12)).grid(row=6, column=0, pady=5, sticky="e")
spin_talla = tk.Spinbox(frame, from_=0, to=300, font=("Helvetica", 12))
spin_talla.grid(row=6, column=1, pady=5)

# Crear y colocar los checkbuttons para seleccionar habilidades profesionales
tk.Label(frame, text="Habilidades", bg="#ffffff", font=("Helvetica", 12)).grid(row=7, column=0, pady=5, sticky="e")
habilidades = ["Liderazgo", "Comunicación", "Trabajo en equipo", "Creatividad", "Resolución de problemas", "Gestión del tiempo", "Adaptabilidad", "Pensamiento crítico", "Negociación", "Planificación"]
habilidades_vars = {habilidad: tk.BooleanVar() for habilidad in habilidades}

for i, (habilidad, var) in enumerate(habilidades_vars.items()):
    tk.Checkbutton(frame, text=habilidad, variable=var, bg="#ffffff", font=("Helvetica", 12)).grid(row=8+i//2, column=i%2, sticky="w")

# Crear un marco para los botones
button_frame = tk.Frame(root, bg="#ffffff")
button_frame.pack(pady=10)

# Crear y colocar los botones para crear PDF, Word y Excel
button_create_pdf = tk.Button(button_frame, text="Crear PDF", command=create_pdf, bg="#800000", fg="#ffffff", font=("Helvetica", 12), padx=10, pady=5)
button_create_pdf.grid(row=0, column=0, padx=5)

button_create_word = tk.Button(button_frame, text="Crear Word", command=create_word, bg="#800000", fg="#ffffff", font=("Helvetica", 12), padx=10, pady=5)
button_create_word.grid(row=0, column=1, padx=5)

button_create_excel = tk.Button(button_frame, text="Crear Excel", command=create_excel, bg="#800000", fg="#ffffff", font=("Helvetica", 12), padx=10, pady=5)
button_create_excel.grid(row=0, column=2, padx=5)

# Ejecutar la aplicación
root.mainloop()




